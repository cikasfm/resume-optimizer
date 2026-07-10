#!/usr/bin/env python3
"""
Generate a beautifully styled DOCX from markdown source, matching
the layout, spacing, and design of the PDF/Print stylesheet.
Uses table-based divider lines to guarantee support in Apple Pages / Quick Look.
"""
import argparse
import re
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config_loader import get_path, get_output_filename

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SRC = get_path('resume_source')
DEFAULT_OUTDIR = get_path('build_dir')
DEFAULT_OUTDIR.mkdir(parents=True, exist_ok=True)
DEFAULT_OUT = DEFAULT_OUTDIR / get_output_filename('baseline_docx')

MARKDOWN_INLINE_RE = re.compile(r'(\*\*|__)(.+?)\1|(\*|_)(.+?)\3')
MARKDOWN_LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')


def parse_markdown_inline(text):
    """Parse inline markdown for bold and italic formatting."""
    segments = []
    last_end = 0

    for match in MARKDOWN_INLINE_RE.finditer(text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], False, False))

        if match.group(1):
            segments.append((match.group(2), True, False))
        else:
            segments.append((match.group(4), False, True))

        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], False, False))

    return segments


def add_hyperlink(paragraph, text, url, color_hex="2B6CB0", size_pt=10.5, underline=True):
    """Adds a clickable hyperlink to a paragraph, explicitly forcing Arial and size."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink node
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r node
    new_run = OxmlElement('w:r')

    # Create a new w:rPr node
    rPr = OxmlElement('w:rPr')

    # Add Arial font styling
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)

    # Add size styling (in half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color_hex)
    rPr.append(c)

    # Add underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    # Add text
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink


def add_basic_runs(paragraph, text, size_pt=10.5, color_rgb=(0x33, 0x33, 0x33), bold_override=None):
    """Adds bold/italic runs to a paragraph, forcing Arial, size, and color."""
    for segment, bold, italic in parse_markdown_inline(text):
        if segment:
            run = paragraph.add_run(segment)
            run.font.name = 'Arial'
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(*color_rgb)
            run.bold = bold_override if bold_override is not None else bold
            run.italic = italic


def add_formatted_runs(paragraph, text, size_pt=10.5, color_rgb=(0x33, 0x33, 0x33), bold_override=None):
    """Parses text for links and basic inline styles and adds them to paragraph."""
    last_end = 0
    for match in MARKDOWN_LINK_RE.finditer(text):
        if match.start() > last_end:
            add_basic_runs(paragraph, text[last_end:match.start()], size_pt, color_rgb, bold_override)
        
        link_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(paragraph, link_text, link_url, size_pt=size_pt)
        last_end = match.end()
        
    if last_end < len(text):
        add_basic_runs(paragraph, text[last_end:], size_pt, color_rgb, bold_override)


def create_divider_heading_table(doc, text, color_hex="1A365D", size_pt=12.5):
    """Creates a single-cell table with a bottom border to represent a heading divider line."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    
    # Remove default table borders
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    table.columns[0].width = docx.shared.Inches(7.0)
    cell = table.cell(0, 0)
    cell.width = docx.shared.Inches(7.0)
    
    # Clear cell margins
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    # Set the bottom border
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    tcBorders.append(bottom)
    
    # Clear other cell borders
    for b in ['top', 'left', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_formatted_runs(p, text.upper(), size_pt=size_pt, color_rgb=(0x1A, 0x36, 0x5D), bold_override=True)
    return table


def create_contact_info_table(doc, text):
    """Creates a single-cell table with a bottom border to represent the contact details block."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    table.columns[0].width = docx.shared.Inches(7.0)
    cell = table.cell(0, 0)
    cell.width = docx.shared.Inches(7.0)
    
    # Clear cell margins
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'bottom', 'left', 'right']:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)
    
    # Set the bottom border
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '8')
    bottom.set(qn('w:color'), 'CCCCCC')
    tcBorders.append(bottom)
    
    # Clear other cell borders
    for b in ['top', 'left', 'right']:
        node = OxmlElement(f'w:{b}')
        node.set(qn('w:val'), 'none')
        tcBorders.append(node)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    add_formatted_runs(p, text, size_pt=9.5, color_rgb=(0x4A, 0x55, 0x68))
    return table


def build_docx(source_path, output_path):
    lines = source_path.read_text(encoding='utf-8').splitlines()
    doc = Document()

    # Set document margins (0.75 inch -> 54pt)
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    h2_count = 0
    is_first_paragraph_after_h1 = False

    for line in lines:
        line = line.rstrip()
        if not line:
            continue
            
        if line.startswith('# '):
            paragraph = doc.add_paragraph('')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            add_formatted_runs(paragraph, line[2:].strip(), size_pt=20, color_rgb=(0x1A, 0x36, 0x5D), bold_override=True)
            is_first_paragraph_after_h1 = True
        elif is_first_paragraph_after_h1:
            # Centered contact info with bottom border divider
            create_contact_info_table(doc, line.strip())
            spacer = doc.add_paragraph('')
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(0)
            is_first_paragraph_after_h1 = False
        elif line.startswith('## '):
            h2_count += 1
            text = line[3:].strip()
            if h2_count == 1:
                # Subtitle (Staff Software Engineer | ...) - no border, smaller font (11pt)
                paragraph = doc.add_paragraph('')
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(10)
                add_formatted_runs(paragraph, text.upper(), size_pt=11, color_rgb=(0x1A, 0x36, 0x5D), bold_override=True)
            else:
                # Section heading (uppercase, bottom border table)
                create_divider_heading_table(doc, text, color_hex="1A365D", size_pt=12.5)
                spacer = doc.add_paragraph('')
                spacer.paragraph_format.space_before = Pt(6)
                spacer.paragraph_format.space_after = Pt(0)
        elif line.startswith('### '):
            paragraph = doc.add_paragraph('')
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            add_formatted_runs(paragraph, line[4:].strip(), size_pt=11, color_rgb=(0x2B, 0x6C, 0xB0), bold_override=True)
        elif line.startswith('#### '):
            paragraph = doc.add_paragraph('')
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(3)
            add_formatted_runs(paragraph, line[5:].strip(), size_pt=10, color_rgb=(0x4A, 0x55, 0x68), bold_override=True)
        elif line.lstrip().startswith('- '):
            paragraph = doc.add_paragraph('', style='List Bullet')
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
            add_formatted_runs(paragraph, line.strip()[2:].strip(), size_pt=10.5, color_rgb=(0x33, 0x33, 0x33))
        else:
            paragraph = doc.add_paragraph('')
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.15
            add_formatted_runs(paragraph, line, size_pt=10.5, color_rgb=(0x33, 0x33, 0x33))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'Wrote {output_path}')


def main():
    parser = argparse.ArgumentParser(description='Generate DOCX from markdown source')
    parser.add_argument('--source', '-s', default=str(DEFAULT_SRC), help='Markdown source file')
    parser.add_argument('--output', '-o', default=str(DEFAULT_OUT), help='Output DOCX file')

    args = parser.parse_args()
    source = Path(args.source).resolve()
    output = Path(args.output).resolve()

    if not source.exists():
        raise FileNotFoundError(f'Markdown source file not found: {source}')

    build_docx(source, output)


if __name__ == '__main__':
    main()
